import json
import logging
import os
import re
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Decomposed document — text blobs + raw PDF bytes for Gemini File API
# ---------------------------------------------------------------------------
@dataclass
class DocumentParts:
    text_parts: list[str] = field(default_factory=list)
    pdf_bytes: list[tuple[str, bytes]] = field(default_factory=list)    # (label, data)
    image_bytes: list[tuple[str, str, bytes]] = field(default_factory=list)  # (label, mime_type, data)


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def parse_pdf(file_path: str) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
            for table in page.extract_tables():
                if not table:
                    continue
                for row in table:
                    cells = [str(c).strip() if c else "" for c in row]
                    parts.append(" | ".join(cells))
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# HTML → text (preserves table/cell structure for Gemini)
# ---------------------------------------------------------------------------
def _html_to_text(html: str) -> str:
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style"]):
            tag.decompose()
        # Keep cell boundaries readable
        for td in soup.find_all(["td", "th"]):
            td.insert_after(" | ")
        for tr in soup.find_all("tr"):
            tr.insert_after("\n")
        return soup.get_text(" ", strip=True)
    except Exception:
        text = re.sub(r"<[^>]+>", " ", html)
        return re.sub(r"\s+", " ", text).strip()


# ---------------------------------------------------------------------------
# Body extractors
# ---------------------------------------------------------------------------
def _msg_body(msg_obj) -> str:
    """Plain-text body for lightweight local scanning (job numbers, mode detection)."""
    body = getattr(msg_obj, "body", None) or ""

    if not body:
        html = getattr(msg_obj, "htmlBody", None)
        if html:
            if isinstance(html, bytes):
                html = html.decode("utf-8", errors="replace")
            body = _html_to_text(html)

    subject = getattr(msg_obj, "subject", None) or ""
    sender  = getattr(msg_obj, "sender",  None) or ""
    header  = "\n".join(filter(None, [
        f"Subject: {subject}" if subject else "",
        f"From: {sender}"    if sender  else "",
    ]))
    return f"{header}\n\n{body}".strip() if header else body.strip()


def _html_tables_to_json(html: str) -> str:
    """
    Convert an HTML email body into structured text:
    - <table> elements become JSON arrays (first row = headers)
    - remaining text is extracted as plain text
    Tables appear inline in document order relative to surrounding text.
    """
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()

    output_parts: list[str] = []

    # Walk top-level children to preserve document order
    for element in soup.body.children if soup.body else soup.children:
        # Recurse into non-table block elements; convert table elements directly
        _collect_parts(element, output_parts)

    return "\n\n".join(p for p in output_parts if p.strip())


def _collect_parts(element, output_parts: list[str]) -> None:
    """Recursively walk element, converting tables to JSON and collecting text."""
    from bs4 import Tag, NavigableString

    if isinstance(element, NavigableString):
        text = element.strip()
        if text:
            output_parts.append(text)
        return

    if not isinstance(element, Tag):
        return

    if element.name == "table":
        table_json = _table_to_json(element)
        if table_json:
            output_parts.append(table_json)
        return

    # For non-table tags: recurse into children
    child_parts: list[str] = []
    for child in element.children:
        _collect_parts(child, child_parts)

    combined = " ".join(child_parts).strip()
    if combined:
        output_parts.append(combined)


def _table_to_json(table_tag) -> str:
    """Convert a BeautifulSoup <table> to a JSON string (first row = headers)."""
    rows_data: list[list[str]] = []
    for tr in table_tag.find_all("tr"):
        cells = [td.get_text(" ", strip=True) for td in tr.find_all(["td", "th"])]
        if any(cells):
            rows_data.append(cells)

    if not rows_data:
        return ""

    headers = rows_data[0]
    result: list[dict] = []
    for row in rows_data[1:]:
        if any(row):
            padded = row + [""] * (len(headers) - len(row))
            result.append(dict(zip(headers, padded)))

    if not result:
        return ""

    return f"[TABLE]\n{json.dumps(result, ensure_ascii=False)}"


def _msg_body_for_ai(msg_obj) -> str:
    """Body for AI extraction: HTML tables become JSON, rest is plain text."""
    subject = getattr(msg_obj, "subject", None) or ""
    sender  = getattr(msg_obj, "sender",  None) or ""
    header  = "\n".join(filter(None, [
        f"Subject: {subject}" if subject else "",
        f"From: {sender}"    if sender  else "",
    ]))

    html = getattr(msg_obj, "htmlBody", None)
    if html:
        if isinstance(html, bytes):
            html = html.decode("utf-8", errors="replace")
        try:
            body = _html_tables_to_json(html)
        except Exception:
            body = getattr(msg_obj, "body", None) or ""
    else:
        body = getattr(msg_obj, "body", None) or ""

    return f"{header}\n\n{body}".strip() if header else body.strip()


# ---------------------------------------------------------------------------
# Attachment walker — returns list of text blobs
# ---------------------------------------------------------------------------
def _walk_attachments(msg_obj) -> list[str]:
    parts: list[str] = []
    attachments = getattr(msg_obj, "attachments", None) or []

    for att in attachments:
        data = getattr(att, "data", None)
        if data is None:
            continue

        # ── Embedded MSGFile (extract-msg MSGAttachment) ──────────────
        # att.data is an MSGFile instance, not bytes
        if hasattr(data, "body") or hasattr(data, "attachments"):
            try:
                nested_body  = _msg_body(data)
                nested_parts = _walk_attachments(data)
                all_p = [p for p in [nested_body] + nested_parts if p.strip()]
                if all_p:
                    parts.append("[Embedded Email]\n" + "\n\n".join(all_p))
            except Exception as exc:
                logger.debug("embedded MSG parse failed: %s", exc)
            continue

        if not isinstance(data, (bytes, bytearray)) or not data:
            continue

        # ── Determine filename ────────────────────────────────────────
        fname = (
            getattr(att, "longFilename",  None)
            or getattr(att, "shortFilename", None)
            or getattr(att, "name",          None)
            or ""
        ).strip()
        ext = Path(fname).suffix.lower() if fname else ""

        # ── PDF attachment ─────────────────────────────────────────────
        if ext == ".pdf" or (not ext and _looks_like_pdf(data)):
            suffix = ".pdf"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            try:
                pdf_text = parse_pdf(tmp_path)
                if pdf_text.strip():
                    label = f"[PDF: {fname}]" if fname else "[PDF Attachment]"
                    parts.append(f"{label}\n{pdf_text}")
            except Exception as exc:
                logger.debug("PDF parse failed (%s): %s", fname, exc)
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

        # ── Nested .msg attachment ─────────────────────────────────────
        elif ext == ".msg":
            with tempfile.NamedTemporaryFile(suffix=".msg", delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            try:
                import extract_msg as _em
                nested     = _em.Message(tmp_path)
                nested_body  = _msg_body(nested)
                nested_parts = _walk_attachments(nested)
                all_p = [p for p in [nested_body] + nested_parts if p.strip()]
                if all_p:
                    label = f"[Attached MSG: {fname}]" if fname else "[Attached MSG]"
                    parts.append(label + "\n" + "\n\n".join(all_p))
            except Exception as exc:
                logger.debug("nested MSG parse failed (%s): %s", fname, exc)
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

        # ── Word attachment (.docx / .doc) ────────────────────────────
        elif ext in _WORD_EXTS | _DOC_EXTS:
            with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            try:
                word_text = parse_docx(tmp_path) if ext in _WORD_EXTS else parse_doc(tmp_path)
                if word_text.strip():
                    label = f"[Word: {fname}]" if fname else "[Word Attachment]"
                    parts.append(f"{label}\n{word_text}")
            except Exception as exc:
                logger.debug("Word parse failed (%s): %s", fname, exc)
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

        # ── Excel attachment ───────────────────────────────────────────
        elif ext in _EXCEL_EXTS | _XLS_EXTS:
            with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            try:
                xl_text = parse_excel(tmp_path)
                if xl_text.strip():
                    label = f"[Excel: {fname}]" if fname else "[Excel Attachment]"
                    parts.append(f"{label}\n{xl_text}")
            except Exception as exc:
                logger.debug("Excel parse failed (%s): %s", fname, exc)
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

    return parts


def _looks_like_pdf(data: bytes) -> bool:
    return data[:4] == b"%PDF"


_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}
_IMAGE_MIME = {
    ".png":  "image/png",
    ".jpg":  "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif":  "image/gif",
    ".webp": "image/webp",
    ".bmp":  "image/bmp",
}
# Images below this threshold are almost certainly logos / spacers / icons.
_IMAGE_MIN_BYTES = 20_000
# Filename fragments that reliably indicate a signature / decorative image.
_SIG_PATTERNS = {"logo", "sig", "banner", "footer", "header", "icon", "avatar", "badge"}


def _looks_like_image(data: bytes) -> tuple[bool, str]:
    """Return (is_image, mime_type) based on magic bytes."""
    if len(data) < 4:
        return False, ""
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        return True, "image/png"
    if data[:3] == b"\xff\xd8\xff":
        return True, "image/jpeg"
    if data[:6] in (b"GIF87a", b"GIF89a"):
        return True, "image/gif"
    if data[:4] == b"RIFF" and len(data) >= 12 and data[8:12] == b"WEBP":
        return True, "image/webp"
    return False, ""


def _is_signature_image(fname: str, data: bytes) -> bool:
    """Return True if this image looks like a logo / signature decoration (skip it)."""
    if len(data) < _IMAGE_MIN_BYTES:
        return True
    return any(pat in fname.lower() for pat in _SIG_PATTERNS)


_EXCEL_EXTS = {".xlsx", ".xlsm", ".xltx", ".xltm"}
_XLS_EXTS   = {".xls", ".xlt"}
_WORD_EXTS  = {".docx", ".docm"}
_DOC_EXTS   = {".doc", ".dot"}


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------
def parse_excel(file_path: str) -> str:
    """Convert all sheets in an Excel workbook to pipe-delimited text."""
    ext = Path(file_path).suffix.lower()
    parts: list[str] = []

    if ext in _EXCEL_EXTS:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows_text: list[str] = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() if c is not None else "" for c in row]
                if any(cells):
                    rows_text.append(" | ".join(cells))
            if rows_text:
                parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows_text))
        wb.close()

    elif ext in _XLS_EXTS:
        try:
            import xlrd
            wb = xlrd.open_workbook(file_path)
            for sheet in wb.sheets():
                rows_text = []
                for rx in range(sheet.nrows):
                    cells = [str(sheet.cell_value(rx, cx)).strip() for cx in range(sheet.ncols)]
                    if any(cells):
                        rows_text.append(" | ".join(cells))
                if rows_text:
                    parts.append(f"[Sheet: {sheet.name}]\n" + "\n".join(rows_text))
        except ImportError:
            raise ValueError(
                "Cannot parse .xls files — install 'xlrd' to enable legacy Excel support."
            )

    return "\n\n".join(parts)


# ---------------------------------------------------------------------------
# Word documents
# ---------------------------------------------------------------------------
def parse_docx(file_path: str) -> str:
    """Extract text from a .docx / .docm file via python-docx."""
    from docx import Document
    doc = Document(file_path)
    parts: list[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_doc(file_path: str) -> str:
    """Best-effort text extraction from a legacy .doc binary file.

    python-docx cannot read the old OLE binary format.  We use olefile
    (already a dependency of extract_msg) to read the WordDocument stream
    and pull out printable text runs.  Results may be incomplete.
    """
    try:
        import olefile
        with olefile.OleFileIO(file_path) as ole:
            if not ole.exists("WordDocument"):
                return ""
            raw = ole.openstream("WordDocument").read()
        # The text in a .doc file is stored as UTF-16-LE runs starting after
        # a large fixed header.  A good-enough heuristic: decode as latin-1
        # and keep sequences of printable ASCII of length >= 4.
        decoded = raw.decode("latin-1", errors="ignore")
        runs = re.findall(r"[ -~\t\r\n]{4,}", decoded)
        return "\n".join(r.strip() for r in runs if r.strip())
    except Exception as exc:
        logger.debug("legacy .doc parse failed (%s): %s", Path(file_path).name, exc)
        return ""


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def parse_msg(file_path: str) -> str:
    import extract_msg
    msg   = extract_msg.Message(file_path)
    body  = _msg_body(msg)
    atts  = _walk_attachments(msg)
    parts = [p for p in [body] + atts if p.strip()]
    result = "\n\n".join(parts)
    logger.debug("parse_msg(%s): body=%d chars, attachments=%d, total=%d chars",
                 Path(file_path).name, len(body), len(atts), len(result))
    return result


def parse_file(file_path: str) -> str:
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(file_path)
    if suffix == ".msg":
        return parse_msg(file_path)
    if suffix in _EXCEL_EXTS | _XLS_EXTS:
        return parse_excel(file_path)
    if suffix in _WORD_EXTS:
        return parse_docx(file_path)
    if suffix in _DOC_EXTS:
        return parse_doc(file_path)
    raise ValueError(f"Unsupported file type: {suffix}")


# ---------------------------------------------------------------------------
# Excel → JSON  (richer structure than pipe-delimited text)
# ---------------------------------------------------------------------------
_EXCEL_MAX_ROWS = 500


def parse_excel_as_json(file_path: str) -> str:
    """Convert an Excel workbook to a JSON string. First non-empty row = headers."""
    ext = Path(file_path).suffix.lower()
    sheets_data = []

    if ext in _EXCEL_EXTS:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            headers: list[str] | None = None
            rows: list[dict] = []
            truncated = False
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() if c is not None else "" for c in row]
                if not any(cells):
                    continue
                if headers is None:
                    headers = cells
                elif len(rows) < _EXCEL_MAX_ROWS:
                    rows.append(dict(zip(headers, cells)))
                else:
                    truncated = True
                    break
            if headers is not None:
                entry: dict = {"sheet": sheet_name, "headers": headers, "rows": rows}
                if truncated:
                    entry["note"] = f"first {_EXCEL_MAX_ROWS} data rows shown; sheet has more rows"
                sheets_data.append(entry)
        wb.close()

    elif ext in _XLS_EXTS:
        try:
            import xlrd
            wb = xlrd.open_workbook(file_path)
            for sheet in wb.sheets():
                headers2: list[str] | None = None
                rows2: list[dict] = []
                truncated2 = False
                for rx in range(sheet.nrows):
                    cells = [str(sheet.cell_value(rx, cx)).strip() for cx in range(sheet.ncols)]
                    if not any(cells):
                        continue
                    if headers2 is None:
                        headers2 = cells
                    elif len(rows2) < _EXCEL_MAX_ROWS:
                        rows2.append(dict(zip(headers2, cells)))
                    else:
                        truncated2 = True
                        break
                if headers2 is not None:
                    entry2: dict = {"sheet": sheet.name, "headers": headers2, "rows": rows2}
                    if truncated2:
                        entry2["note"] = f"first {_EXCEL_MAX_ROWS} data rows shown; sheet has more rows"
                    sheets_data.append(entry2)
        except ImportError:
            raise ValueError("Cannot parse .xls files — install 'xlrd'.")

    return json.dumps(sheets_data, ensure_ascii=False)


# ---------------------------------------------------------------------------
# Decompose walker — collects PDF bytes + text from MSG attachments
# ---------------------------------------------------------------------------
def _walk_for_decompose(msg_obj, parts: DocumentParts) -> None:
    for att in getattr(msg_obj, "attachments", None) or []:
        data = getattr(att, "data", None)
        if data is None:
            continue

        # Embedded MSG
        if hasattr(data, "body") or hasattr(data, "attachments"):
            try:
                nested_body = _msg_body_for_ai(data)
                if nested_body.strip():
                    parts.text_parts.append(f"[Embedded Email]\n{nested_body}")
                _walk_for_decompose(data, parts)
            except Exception as exc:
                logger.debug("embedded MSG decompose failed: %s", exc)
            continue

        if not isinstance(data, (bytes, bytearray)) or not data:
            continue

        fname = (
            getattr(att, "longFilename", None)
            or getattr(att, "shortFilename", None)
            or getattr(att, "name", None)
            or ""
        ).strip()
        ext = Path(fname).suffix.lower() if fname else ""

        if ext == ".pdf" or (not ext and _looks_like_pdf(data)):
            parts.pdf_bytes.append((fname or "attachment.pdf", bytes(data)))

        elif ext == ".msg":
            with tempfile.NamedTemporaryFile(suffix=".msg", delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            try:
                import extract_msg as _em
                nested = _em.Message(tmp_path)
                nested_body = _msg_body_for_ai(nested)
                if nested_body.strip():
                    parts.text_parts.append(f"[Attached MSG: {fname or 'nested.msg'}]\n{nested_body}")
                _walk_for_decompose(nested, parts)
            except Exception as exc:
                logger.debug("nested MSG decompose failed (%s): %s", fname, exc)
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

        elif ext in _EXCEL_EXTS | _XLS_EXTS:
            with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            try:
                xl_json = parse_excel_as_json(tmp_path)
                parts.text_parts.append(f"[Excel: {fname or 'attachment'}]\n{xl_json}")
            except Exception as exc:
                logger.debug("Excel decompose failed (%s): %s", fname, exc)
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

        elif ext in _WORD_EXTS | _DOC_EXTS:
            with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            try:
                word_text = parse_docx(tmp_path) if ext in _WORD_EXTS else parse_doc(tmp_path)
                if word_text.strip():
                    parts.text_parts.append(f"[Word: {fname or 'attachment'}]\n{word_text}")
            except Exception as exc:
                logger.debug("Word decompose failed (%s): %s", fname, exc)
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

        elif ext in _IMAGE_EXTS or _looks_like_image(data)[0]:
            # Determine MIME type from extension, or fall back to magic bytes
            mime = _IMAGE_MIME.get(ext) or _looks_like_image(data)[1]
            if mime and not _is_signature_image(fname, data):
                label = fname or f"image{len(parts.image_bytes) + 1}{ext or '.img'}"
                parts.image_bytes.append((label, mime, bytes(data)))
                logger.debug("body image queued for upload: %s (%d bytes)", label, len(data))
            else:
                logger.debug("signature/small image skipped: %s (%d bytes)", fname, len(data))


# ---------------------------------------------------------------------------
# Public decompose API
# ---------------------------------------------------------------------------
def decompose_file(file_path: str) -> DocumentParts:
    """Decompose any supported file into text parts + PDF bytes for Gemini File API."""
    parts = DocumentParts()
    suffix = Path(file_path).suffix.lower()

    if suffix == ".pdf":
        parts.pdf_bytes.append((Path(file_path).name, Path(file_path).read_bytes()))

    elif suffix == ".msg":
        import extract_msg
        msg = extract_msg.Message(file_path)
        body = _msg_body_for_ai(msg)
        if body.strip():
            parts.text_parts.append(body)
        _walk_for_decompose(msg, parts)

    elif suffix in _EXCEL_EXTS | _XLS_EXTS:
        parts.text_parts.append(parse_excel_as_json(file_path))

    elif suffix in _WORD_EXTS:
        text = parse_docx(file_path)
        if text.strip():
            parts.text_parts.append(text)

    elif suffix in _DOC_EXTS:
        text = parse_doc(file_path)
        if text.strip():
            parts.text_parts.append(text)

    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    return parts
